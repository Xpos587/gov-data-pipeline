import re
import logging
from typing import Optional, List, Dict, Any, Set
from io import BytesIO

import aiohttp
import polars as pl
from docx import Document

from .base import BaseHandler
from utils.gpt import process_table
from utils.pdf2docx import convert_to_docx

logger = logging.getLogger(__name__)


async def pdf_to_dataframe(
    pdf_content: bytes,
    user_agent: str,
    session: aiohttp.ClientSession,
    proxy: Optional[str] = None,
    proxy_auth: Optional[aiohttp.BasicAuth] = None,
) -> Optional[pl.DataFrame]:
    """
    Конвертирует PDF в DOCX с помощью convert_to_docx, затем извлекает таблицы из DOCX
    и формирует pl.DataFrame.
    """
    docx_data = await convert_to_docx(
        pdf_content, user_agent, session, proxy, proxy_auth
    )
    if not docx_data:
        logger.error("Ошибка конвертации PDF в DOCX.")
        return None

    try:
        document = Document(BytesIO(docx_data))
        all_data: List[List[str]] = []
        for table in document.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                all_data.append(row_data)
        if not all_data:
            logger.error("Таблицы не найдены в DOCX.")
            return None

        # Выравниваем длину строк и формируем заголовки
        max_len = max(len(row) for row in all_data)
        padded = [row + [""] * (max_len - len(row)) for row in all_data]
        raw_columns = padded[0]
        columns: List[str] = [
            re.sub(r"\s+", " ", col).strip() or f"Unnamed_{i}"
            for i, col in enumerate(raw_columns)
        ]
        if len(padded) < 2:
            logger.error("Недостаточно строк для формирования таблицы.")
            return None

        data_rows = padded[1:]
        df = pl.DataFrame(data_rows, schema=columns, orient="row")

        # Функция для очистки содержимого ячеек
        def clean_cell(cell: Optional[str]) -> str:
            return re.sub(r"\s+", " ", cell.strip()) if cell else ""

        # Используем map_elements + cast(Utf8) для корректного определения типа
        df = df.with_columns(
            [
                pl.col(col).map_elements(clean_cell, return_dtype=pl.Utf8).alias(col)
                for col in columns
            ]
        )

        return df
    except Exception as e:
        logger.error(f"Ошибка при обработке DOCX: {e}")
        return None


class KyrgyzstanHandler(BaseHandler):
    """
    Асинхронный хендлер для получения данных с сайта Киргизии.
    """

    IMAGE_COLUMN_NAME: str = "Наименование (вид, описание, изображение) ОИС"
    TRADEMARK_COLUMN_NAME: str = "Наименование (вид, описание, изображение) ОИС"
    DESCRIPTION_COLUMN_NAME: str = "Наименование товаров, в отношении которых принимаются меры Класс товаров по МКТУ/Код товаров по ТНВЭД"
    ROW_OFFSET = 0

    async def retrieve(
        self,
        user_agent: Optional[str],
        proxy: Optional[str] = None,
        proxy_auth: Optional[aiohttp.BasicAuth] = None,
    ) -> Optional[bytes]:
        """
        Загружает страницу с PDF, извлекает ссылку на PDF и возвращает его содержимое.
        """
        page_url: str = "https://www.customs.gov.kg/site/ru/master/customskg/intellektualdyk-menchik-ukuktaryn-korgoo"
        logger.info(f"Запрашиваем страницу с PDF: {page_url}")
        pdf_page = await self.fetch(
            url=page_url,
            proxy=proxy,
            proxy_auth=proxy_auth,
            user_agent=user_agent,
        )
        if not pdf_page:
            logger.error("Не удалось загрузить страницу с PDF.")
            return None

        try:
            page_text = pdf_page.decode("utf-8")
        except UnicodeDecodeError as e:
            logger.error(f"Ошибка декодирования страницы: {e}")
            return None

        pdf_match = re.search(
            r'<a\s+href="([^"]+\.pdf)"[^>]*>.*?Таможенный\s+реестр.*?интеллектуальной\s+собственности.*?</a>',
            page_text,
            re.IGNORECASE,
        )
        if not pdf_match:
            logger.warning("Ссылка на PDF не найдена на странице.")
            return None

        pdf_path = pdf_match.group(1)
        pdf_url = f"https://www.customs.gov.kg{pdf_path}"
        logger.info(f"Найдена ссылка на PDF: {pdf_url}")

        pdf_content = await self.fetch(
            url=pdf_url,
            proxy=proxy,
            proxy_auth=proxy_auth,
            user_agent=user_agent,
        )
        if not pdf_content:
            logger.error("Не удалось скачать PDF.")
            return None

        return pdf_content

    async def process(
        self,
        user_agent: str,
        proxy: Optional[str] = None,
        proxy_auth: Optional[aiohttp.BasicAuth] = None,
        correction: bool = False,
    ) -> Optional[pl.DataFrame]:
        """
        Получает PDF, конвертирует его в DOCX, извлекает таблицы и предобрабатывает их по методике из ipynb.
        """
        pdf_content = await self.retrieve(user_agent, proxy, proxy_auth)
        if not pdf_content:
            logger.error("PDF не получен, прерываем обработку.")
            return None

        if self.session is None:
            raise RuntimeError(
                "Сессия не инициализирована. Используйте 'async with' для работы с хендлером."
            )

        # Конвертация PDF в DOCX
        docx_data = await convert_to_docx(
            pdf_content, user_agent, self.session, proxy, proxy_auth
        )
        if not docx_data:
            logger.error("Ошибка конвертации PDF в DOCX.")
            return None

        def clean_cell(cell: Optional[str]) -> str:
            """
            Удаляем лишние пробелы, переводы строк и т.п.
            Если ячейка None – возвращаем пустую строку.
            """
            return re.sub(r"\s+", " ", cell.strip()) if cell else ""

        def is_new_record(val: str) -> bool:
            """
            Проверяем, выглядит ли значение как новый номер записи (например, '№0001/ТЗ').
            """
            pattern = r"^(?:№?\d{4,})(/ТЗ.*)?"
            return bool(re.match(pattern, val.strip()))

        def merge_continued_rows(df: pl.DataFrame, key_col: str) -> pl.DataFrame:
            """
            Склеивает «продолжения» строк. Если значение в key_col не соответствует формату «новая запись», то объединяем с предыдущей.
            """
            records: List[Dict[str, Any]] = df.to_dicts()
            merged_rows: List[Dict[str, Any]] = []
            prev: Optional[Dict[str, Any]] = None

            for row in records:
                current_val = str(row[key_col]).strip()
                if current_val.startswith("Name:"):
                    continue

                if is_new_record(current_val):
                    if prev is not None:
                        merged_rows.append(prev)
                    prev = row
                else:
                    if prev is not None:
                        for col in df.columns:
                            cur_val = str(row[col]).strip()
                            if cur_val:
                                old_val = str(prev.get(col, "")).strip()
                                prev[col] = (
                                    (old_val + " " + cur_val).strip()
                                    if old_val
                                    else cur_val
                                )
                    else:
                        prev = row

            if prev is not None:
                merged_rows.append(prev)

            return pl.DataFrame(merged_rows)

        def preprocess_reg_num(value: str) -> str:
            """
            Нормализует значение столбца 'Рег. №' по заданным правилам.
            """
            val = value.strip()
            if not val:
                return val
            val = re.sub(r"^№\s*", "", val)  # Убираем ведущий "№"
            val = re.sub(r"\s*См\.\s*", " См. ", val)  # Стабилизируем "См."
            val = re.sub(r"\s*[–—]\s*", "-", val)  # Длинные дефисы → "-"
            val = re.sub(r"\s*-\s*", "-", val)
            val = re.sub(r"\s*/\s*", "/", val)
            val = re.sub(r"\s*\.\s*", ".", val)
            val = re.sub(r"(?<=\d)\s+(?=\d)", "", val)  # Убираем пробелы между цифрами
            val = re.sub(r"(См\.)\s*(?=\S)", r"\1 ", val)
            val = re.sub(r"-{2,}", "-", val)  # Многократные дефисы
            val = re.sub(r"\s+", " ", val)  # Лишние пробелы
            return val.strip()

        try:
            document = Document(BytesIO(docx_data))
            all_data: List[List[str]] = []
            for table in document.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    all_data.append(row_data)
            if not all_data:
                logger.error("Таблицы не найдены в DOCX.")
                return None

            # Шаг 1. Выравниваем строки по максимальной длине
            max_len = max(len(row) for row in all_data)
            padded_data = [row + [""] * (max_len - len(row)) for row in all_data]

            # Шаг 2. Обрабатываем имена столбцов (первая строка)
            column_names = padded_data[0]
            unique_columns: List[str] = []
            seen: Set[str] = set()
            for col in column_names:
                if not col or col in seen:
                    counter = 1
                    new_col = f"{col or 'Unnamed'}_{counter}"
                    while new_col in seen:
                        counter += 1
                        new_col = f"{col or 'Unnamed'}_{counter}"
                    unique_columns.append(new_col)
                else:
                    unique_columns.append(col.strip())
                seen.add(unique_columns[-1])

            # Создаём DataFrame, пропуская первые 2 строки
            df = pl.DataFrame(padded_data[2:], schema=unique_columns, orient="row")

            # Приводим все столбцы к строковому типу и чистим каждую ячейку
            df = df.with_columns(
                [
                    pl.col(col)
                    .cast(pl.Utf8)
                    .map_elements(clean_cell, return_dtype=pl.Utf8)
                    .alias(col)
                    for col in df.columns
                ]
            )

            # Переименование столбцов (при необходимости)
            df = df.rename(
                {
                    "Рег. №": "Рег. №",
                    "Наименование (вид, описание, изображение) ОИС": "Наименование (вид, описание, изображение) ОИС",
                    "Наименова ние, №, дата документа об охраноспос обности ОИС": "Наименование, №, дата документа об охраноспособности ОИС",
                    "Наименование товаров, в отношении которых принимаются меры Класс товаров по МКТУ/Код товаров по ТНВЭД": "Наименование товаров, в отношении которых принимаются меры (класс товаров по МКТУ/Код товаров по ТНВЭД)",
                    "Правообладате ль": "Правообладатель",
                    "Доверенные лица правообладателя": "Доверенные лица правообладателя",
                    "Срок несения ОИС в Реестр": "Срок внесения ОИС",
                    "Номер и дата письма ГТС": "Номер и дата письма ГТС",
                },
                strict=False,
            )

            # Предобработка столбца "Рег. №"
            if "Рег. №" in df.columns:
                df = df.with_columns(
                    pl.col("Рег. №")
                    .map_elements(preprocess_reg_num, return_dtype=pl.Utf8)
                    .alias("Рег. №")
                )
            else:
                logger.warning("Внимание: в DataFrame нет столбца 'Рег. №'!")

            # Объединяем продолжения строк по "Рег. №"
            df = merge_continued_rows(df, key_col="Рег. №")

        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX: {e}")
            return None

        # Предобработка изображений из исходного PDF
        byte_stream = BytesIO(pdf_content)
        df = await self.process_excel_images(df, byte_stream)

        logger.info(f"Данные успешно загружены и предобработаны: {df.shape}")

        df = await process_table(
            df,
            brand_column=self.TRADEMARK_COLUMN_NAME,
            description_column=self.DESCRIPTION_COLUMN_NAME,
            correction=correction,
        )

        logger.info(f"Формирование таблицы завершено: {df.shape}")
        return df
